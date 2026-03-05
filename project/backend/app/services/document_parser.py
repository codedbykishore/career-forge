"""
Document Parser Service
=======================
Extracts text from various document formats.
"""

from typing import Optional, Dict, Any, List
from pathlib import Path
import hashlib
import structlog

from app.services.bedrock_client import bedrock_client


logger = structlog.get_logger()


class DocumentParserService:
    """
    Service for parsing and extracting text from documents.
    Supports PDF, DOCX, TXT, and Markdown files.
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
    
    def __init__(self):
        pass
    
    def compute_file_hash(self, content: bytes) -> str:
        """Compute SHA-256 hash of file content for deduplication."""
        return hashlib.sha256(content).hexdigest()
    
    async def extract_text(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the file
            file_content: Optional file content bytes (if already loaded)
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext == ".pdf":
            return await self._extract_from_pdf(file_path, file_content)
        elif ext == ".docx":
            return await self._extract_from_docx(file_path, file_content)
        elif ext in {".txt", ".md"}:
            return await self._extract_from_text(file_path, file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    async def _extract_from_pdf(self, file_path: str, content: Optional[bytes] = None) -> str:
        """Extract text from PDF using PyMuPDF."""
        import fitz  # PyMuPDF
        
        if content:
            doc = fitz.open(stream=content, filetype="pdf")
        else:
            doc = fitz.open(file_path)
        
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        
        doc.close()
        return "\n".join(text_parts)
    
    async def _extract_from_docx(self, file_path: str, content: Optional[bytes] = None) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        import io
        
        if content:
            doc = Document(io.BytesIO(content))
        else:
            doc = Document(file_path)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    async def _extract_from_text(self, file_path: str, content: Optional[bytes] = None) -> str:
        """Extract text from plain text or markdown files."""
        if content:
            return content.decode("utf-8")
        
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    async def classify_document(self, text: str) -> Dict[str, Any]:
        """
        Classify document type and extract metadata using Gemini.
        
        Args:
            text: Extracted document text
            
        Returns:
            Dict with doc_type and extracted metadata
        """
        prompt = f"""Analyze this document and classify it.

DOCUMENT TEXT (first 3000 chars):
{text[:3000]}

Classify the document and extract relevant information.
Return a JSON object with:
- "doc_type": one of ["resume", "cover_letter", "project", "certificate", "reference", "other"]
- "confidence": 0.0 to 1.0
- "metadata": relevant extracted fields based on doc_type

For resumes, extract: name, email, phone, skills, education, experience_years
For certificates: issuer, title, date, credential_id
For projects: title, description, technologies
For other: summary

Return ONLY valid JSON."""

        try:
            result = await bedrock_client.generate_json(
                prompt=prompt,
                system_instruction="You are a document classifier. Analyze documents and extract structured information accurately.",
                temperature=0.1,
            )
            return result
        except Exception as e:
            logger.error(f"Document classification failed: {e}")
            return {
                "doc_type": "other",
                "confidence": 0.0,
                "metadata": {}
            }
    
    async def extract_projects_from_resume(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract project information from a resume document.
        
        Args:
            text: Resume text content
            
        Returns:
            List of extracted project dicts
        """
        prompt = f"""Extract all projects mentioned in this resume.

RESUME TEXT:
{text[:5000]}

For each project found, extract:
- title: Project name
- description: Brief description
- technologies: List of technologies/tools used
- highlights: List of achievements/bullet points
- dates: Start and end dates if mentioned

Return a JSON array of project objects. Only include information explicitly stated.
If no projects found, return empty array [].

Return ONLY valid JSON array."""

        try:
            result = await bedrock_client.generate_json(
                prompt=prompt,
                system_instruction="You are a resume parser. Extract project information accurately. Never invent information not present in the text.",
                temperature=0.1,
            )
            if isinstance(result, list):
                return result
        except Exception as e:
            logger.error(f"Project extraction failed: {e}")
        
        return []
    
    async def extract_skills_from_text(self, text: str) -> List[str]:
        """
        Extract skills/technologies from document text.
        
        Args:
            text: Document text
            
        Returns:
            List of extracted skills
        """
        prompt = f"""Extract all technical skills, technologies, frameworks, and tools mentioned in this text.

TEXT:
{text[:4000]}

Return a JSON array of skill/technology names.
Only include specific technical skills (e.g., "Python", "React", "AWS", "Docker").
Do not include soft skills.

Return ONLY a JSON array like: ["Python", "JavaScript", "AWS"]"""

        try:
            result = await bedrock_client.generate_json(
                prompt=prompt,
                temperature=0.1,
            )
            if isinstance(result, list):
                return result
        except Exception as e:
            logger.error(f"Skill extraction failed: {e}")
        
        return []


# Global instance
document_parser = DocumentParserService()
